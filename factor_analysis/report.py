from docx import Document
from docx.shared import Inches, Cm
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Inches, Cm
from docx.oxml.ns import qn
from utils.doc_tool import add_heading, add_table, doc_add_dataframe, doc_add_images


def create_pool_doc(param):
    """
    Generate a DOC file and specify the rules
    :param param:
    :param doc:
    :param level_num:
    :param seq_num:
    :return:
    """

    # 生成文档
    doc = Document()
    default_section = doc.sections[0]
    default_section.page_width = Cm(
        30
    )  # Change the paper size to custom to accommodate large tables and images
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    add_heading("Factor inclusion", doc, level=0, seq=1)

    add_heading("Factor library screening rules", doc, level=1, seq=1)
    Barra_rule = param.pool_thresh["Barra"]
    tech_rule = param.pool_thresh["技术面"]
    fin_rule = param.pool_thresh["基本面"]
    doc.add_paragraph(
        "1. Neutralization rules:\n"
        "        Barra factors are only neutralized for industry and market capitalization;\n"
        "        Both technical indicators and fundamental factors are neutralized for industry, market capitalization, and Barra factors，"
        "Among them, Barra neutralization is applied to Barra factors with a correlation above 0.3\n"
    )
    doc.add_paragraph(
        "2. Correlation screening rules:\n"
        "        Factors with a correlation greater than 0.9 are considered identical and removed, while the factor with the higher long-short backtest Sharpe ratio is retained.\n"
    )
    doc.add_paragraph(
        "3. Factor indicator screening rules: \n"
        "        Barra factors: The mean RankIC must be greater than %s, the t-value must be greater than %s, and the directional consistency probability must be greater than %s."
        "The annualized return of the factor long-short portfolio must be greater than %s, the t-value must be greater than %s, the Sharpe ratio must be greater than %s, and the win rate must be greater than %s.\n"
        "        Technical factors: The mean RankIC must be greater than %s, the t-value must be greater than %s, and the directional consistency probability must be greater than %s."
        "The annualized return of the factor long-short portfolio must be greater than %s, the t-value must be greater than %s, the Sharpe ratio must be greater than %s, and the win rate must be greater than %s.\n"
        "        Fundamental factors: The mean RankIC must be greater than %s, the t-value must be greater than %s, and the directional consistency probability must be greater than %s."
        "The annualized return of the factor long-short portfolio must be greater than %s, the t-value must be greater than %s, the Sharpe ratio must be greater than %s, and the win rate must be greater than %s.\n"
        % (
            Barra_rule["IC"],
            Barra_rule["t_value"],
            Barra_rule["direction_rate"],
            Barra_rule["ARR"],
            Barra_rule["t_value"],
            Barra_rule["sharp"],
            Barra_rule["win_rate"],
            tech_rule["IC"],
            tech_rule["t_value"],
            tech_rule["direction_rate"],
            tech_rule["ARR"],
            tech_rule["t_value"],
            tech_rule["sharp"],
            tech_rule["win_rate"],
            fin_rule["IC"],
            fin_rule["t_value"],
            fin_rule["direction_rate"],
            fin_rule["ARR"],
            fin_rule["t_value"],
            fin_rule["sharp"],
            fin_rule["win_rate"],
        )
    )
    return doc


def add_heading_nonum(title, doc, level, seq):
    if level == 0:
        doc.add_heading(title, level=level)
    else:
        if seq == 0:
            doc.add_heading(title, level=level)
        else:
            doc.add_heading("%s、" % seq + title, level=level)


# def make_report(doc, factor_neut):
#     add_heading("Neutralized factor correlation", doc, level=1, seq=2)
#     chosen_col = [col_ for col_ in factor_neut.columns if "neut" in col_]
#     neut_corr = factor_neut[chosen_col].corr()
#     add_heading_nonum("中性化后因子相关性热力图", doc, level=2, seq=0)
#     draw_thermodynamic_diagram(neut_corr, doc, "相关性热力图", inch=10.5)

#     # 输出结果，先输出总表，再逐个因子单独
#     style = "Light List Accent 2"
#     add_heading("因子表现", doc, level=1, seq=2)
#     barra_df = result_df[result_df["因子"]["因子名"].isin(param.barra_factor)]
#     tech_df = result_df[result_df["因子"]["因子名"].isin(param.tech_facotr)]
#     fin_df = result_df[result_df["因子"]["因子名"].isin(param.fin_factor)]
#     add_heading_nonum("Barra因子表现", doc, level=2, seq=1)
#     add_table(barra_df.T.reset_index().T, doc, size=7, style=style, width=8)
#     add_heading_nonum("技术因子表现", doc, level=2, seq=2)
#     add_table(tech_df.T.reset_index().T, doc, size=7, style=style, width=8)
#     add_heading_nonum("基本面因子表现", doc, level=2, seq=3)
#     add_table(fin_df.T.reset_index().T, doc, size=7, style=style, width=8)
#     # result_df.to_excel(os.path.join(param.pool_test, '因子监控.xlsx'))
#     # result_df = pd.read_excel(os.path.join(param.pool_test, '因子监控.xlsx'))

#     # 3、因子筛选结果
#     add_heading("因子筛选结果", doc, level=1, seq=3)
#     factor_success = pd.DataFrame(columns=result_df.columns)
#     factor_fail = pd.DataFrame(columns=["因子", "分类", "未通过原因"])

#     # 进行因子表现剔除
#     for factor in param.factor_list:
#         factor_df = result_df[result_df["因子"]["因子名"] == factor]
#         factor_cate = factor_df["因子"]["分类"].iloc[0]
#         factor_cate = "技术面"  ############################## 删
#         IC_thresh = param.pool_thresh[factor_cate]["IC"]
#         t_value_thresh = param.pool_thresh[factor_cate]["t_value"]
#         direction_rate_thresh = param.pool_thresh[factor_cate]["direction_rate"]
#         ARR_thresh = param.pool_thresh[factor_cate]["ARR"]
#         win_rate_thresh = param.pool_thresh[factor_cate]["win_rate"]
#         sharp_thresh = param.pool_thresh[factor_cate]["sharp"]

#         factor_word = ""
#         if float(factor_df["RankIC"]["均值"].iloc[0][:-1]) < IC_thresh * 100:
#             factor_word += "因子RankIC的均值过低；"
#         if float(factor_df["RankIC"]["t值"].iloc[0]) < t_value_thresh:
#             factor_word += "因子RankIC的t值过低；"
#         if (
#             float(factor_df["RankIC"]["方向延续概率"].iloc[0][:-1])
#             < direction_rate_thresh * 100
#         ):
#             factor_word += "因子RankIC的方向延续概率过低；"
#         if float(factor_df["多空组合"]["年化收益"].iloc[0][:-1]) < ARR_thresh * 100:
#             factor_word += "因子多空组合收益率过低；"
#         if float(factor_df["多空组合"]["收益率t值"].iloc[0]) < t_value_thresh:
#             factor_word += "因子多空收益率的t值过低；"
#         if float(factor_df["多空组合"]["夏普比"].iloc[0]) < sharp_thresh:
#             factor_word += "因子多空组合夏普比过低；"
#         if float(factor_df["多空组合"]["胜率"].iloc[0][:-1]) < win_rate_thresh * 100:
#             factor_word += "因子多空组合胜率过低；"

#         if factor_word == "":  # 因子通过筛选
#             factor_success = pd.concat([factor_success, factor_df], axis=0)
#         else:
#             factor_fail.loc[len(factor_fail)] = [factor, factor_cate, factor_word]

#     # 进行相关性剔除
#     success_list = list(factor_success.index)
#     neut_corr_success = neut_corr.copy()
#     col_name = [col_[:-5] for col_ in neut_corr_success.columns]
#     neut_corr_success.columns = col_name
#     neut_corr_success.index = col_name

#     neut_corr_success = neut_corr_success[success_list].T[success_list]
#     delete_factor = []
#     for factor_ in success_list:
#         factor_corr = neut_corr_success[factor_]
#         try:
#             factor_del_list = list(factor_corr[abs(factor_corr) > 0.9].index)
#             factor_del_list.append(factor_)
#             del_df = result_df[result_df.index.isin(factor_del_list)]
#             del_df[("多空组合", "夏普比")] = del_df.apply(
#                 lambda x: float(x[("多空组合", "夏普比")]), axis=1
#             )
#             del_df.sort_values(
#                 by=[("多空组合", "夏普比")], ascending=False, inplace=True
#             )  # 按夏普比排序
#             delete_factor += list(del_df.index)[1:]  # 只保存夏普最高的因子
#         except:
#             pass
#     if len(delete_factor) != 0:  # 存在同质化因子
#         delete_factor = set(delete_factor)  # 去重
#         for factor_ in delete_factor:
#             factor_success = factor_success[
#                 factor_success.index != factor_
#             ]  # 从入选中剔除
#             factor_cate = result_df[result_df["因子"]["因子名"] == factor_]["因子"][
#                 "分类"
#             ].iloc[0]
#             factor_word = "与已入选因子相关性过高"
#             factor_fail.loc[len(factor_fail)] = [factor_, factor_cate, factor_word]

#     add_heading("入选因子表现", doc, level=3, seq=1)
#     add_table(factor_success.T.reset_index().T, doc, size=7, style=style, width=8)
#     add_heading("未入选因子落选原因", doc, level=3, seq=2)
#     add_table(factor_fail.T.reset_index().T, doc, size=7, style=style, width=8)

#     # 4、逐个因子表现输出
#     add_heading("逐个因子表现", doc, level=1, seq=4)
#     num = 0
#     for factor in param.factor_list:
#         num += 1
#         add_heading_nonum("%s因子表现" % factor, doc, level=2, seq=num)
#         # 因子收益回测
#         factor_result = result_df[result_df["因子"]["因子名"] == factor]
#         add_heading("收益回测", doc, level=3, seq=1)
#         add_table(factor_result.T.reset_index().T, doc, size=7, style=style, width=8)
#         # 因子相关性
#         add_heading("因子TOP10相关性", doc, level=3, seq=2)
#         factor_corr = neut_corr[[factor + "_neut"]]
#         factor_corr["abs"] = abs(factor_corr[factor + "_neut"])
#         factor_corr.index = [factor_[:-5] for factor_ in factor_corr.index]
#         factor_corr = factor_corr[factor_corr["abs"] != 1]  # 删除因子本身
#         factor_corr.sort_values(by="abs", ascending=False, inplace=True)
#         factor_corr = factor_corr.iloc[:10, :]  # 取前十相关因子
#         factor_corr = round(factor_corr, 2)  # 保留两位小数
#         del factor_corr["abs"]
#         add_table(
#             factor_corr.reset_index().T.reset_index(), doc, size=7, style=style, width=8
#         )
#         # 因子回测曲线
#         add_heading("因子回测曲线", doc, level=3, seq=3)
#         images = os.path.join(
#             param.pool_test + "%s_净值曲线及回撤（多空收益）.png" % factor
#         )
#         doc_add_images(doc, images, inch=6)
#         images = os.path.join(
#             param.pool_test + "%s_净值曲线及回撤（相对收益）.png" % factor
#         )
#         doc_add_images(doc, images, inch=6)

#     doc.save("因子池筛选_磐松.docx")  # 保存路径
