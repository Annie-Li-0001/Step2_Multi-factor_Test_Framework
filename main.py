"""

"""

import os
import time
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Cm
from docx.oxml.ns import qn  # 设置字体。
from Config.job_config import FactorPara
import seaborn as sns

from utils.doc_tool import add_heading, add_table, doc_add_dataframe, doc_add_images

from factor_analysis.plot import draw_thermodynamic_diagram
from factor_analysis.analyse import factor_pool_test_simple

from factor_analysis.data_io import (
    get_calendar_data,
    get_factor_data,
    get_stock_industry,
    get_stock_return,
)


def add_heading_nonum(title, doc, level, seq):
    if level == 0:
        doc.add_heading(title, level=level)
    else:
        if seq == 0:
            doc.add_heading(title, level=level)
        else:
            doc.add_heading("%s、" % seq + title, level=level)


def create_pool_doc(param):
    """
    生成doc文件并标明规则
    :param param:
    :param doc:
    :param level_num:
    :param seq_num:
    :return:
    """

    # 生成文档
    doc = Document()
    default_section = doc.sections[0]
    default_section.page_width = Cm(30)  # 纸张大小改为自定义，方便放下大表和大图
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    add_heading("Neutralized factor correlation", doc, level=0, seq=1)
    # 增加因子库筛选规则说明
    add_heading("Factor library screening rules", doc, level=1, seq=1)
    Barra_rule = param.pool_thresh["Barra"]
    tech_rule = param.pool_thresh["技术面"]
    fin_rule = param.pool_thresh["基本面"]
    doc.add_paragraph(
        "1.Neutralization rules:\n"
        "        Barra factors are only neutralized for industry and market capitalization;\n"
        "        Both technical indicators and fundamental factors are neutralized for industry, market capitalization, and Barra factors,"
        "Among them, Barra neutralization is applied to Barra factors with a correlation above 0.3\n"
    )
    doc.add_paragraph(
        "2. Correlation screening rules:\n"
        "        Factors with a correlation greater than 0.9 are considered identical and removed, while the factor with the higher long-short backtest Sharpe ratio is retained.\n"
    )
    doc.add_paragraph(
        "3.Factor indicator screening rules:\n"
        "        Barra factors: The mean RankIC must be greater than %s, the t-value must be greater than %s, and the directional consistency probability must be greater than %s."
        "The annualized return of the factor long-short portfolio must be greater than %s, the t-value must be greater than %s, the Sharpe ratio must be greater than %s, and the win rate must be greater than %s.\n"
        "        Technical factors: The mean RankIC must be greater than %s, the t-value must be greater than %s, and the directional consistency probability must be greater than %s."
        "The annualized return of the factor long-short portfolio must be greater than %s, the t-value must be greater than %s, the Sharpe ratio must be greater than %s, and the win rate must be greater than %s.\n"
        "        Fundamental factors: The mean RankIC must be greater than %s, the t-value must be greater than %s, and the directional consistency probability must be greater than %s."
        "The annualized return of the factor long-short portfolio must be greater than %s, the t-value must be greater than %s, the Sharpe ratio must be greater than %s, and the win rate must be greater than %s.\n"
        % (
            Barra_rule["IC"],
            Barra_rule["t_value"],
            Barra_rule["direction_rate"],
            Barra_rule["ARR"],
            Barra_rule["t_value"],
            Barra_rule["sharp"],
            Barra_rule["win_rate"],
            tech_rule["IC"],
            tech_rule["t_value"],
            tech_rule["direction_rate"],
            tech_rule["ARR"],
            tech_rule["t_value"],
            tech_rule["sharp"],
            tech_rule["win_rate"],
            fin_rule["IC"],
            fin_rule["t_value"],
            fin_rule["direction_rate"],
            fin_rule["ARR"],
            fin_rule["t_value"],
            fin_rule["sharp"],
            fin_rule["win_rate"],
        )
    )
    return doc


def factor_test() -> pd.DataFrame:
    # (factor_name, factor_neut, trade_days, stock_return,
    #  industry_code_df, drop_industry, index_code, param)
    for factor_name_ in factor_neut.columns[:]:
        factor_values = factor_neut[[factor_name_]]
        df_neut = factor_neut[
            ["date", "next_date", "code", "stock_return", "%s_neut" % factor]
        ]
    df = pd.DataFrame()
    return df


if __name__ == "__main__":

    param = FactorPara()

    print("0. Create a document")
    doc = create_pool_doc(param)

    a = time.time()
    print("1.Data acquisition, neutralization, and correlation calculation.")
    print("1.1 Load the factor data")
    factor_data = get_factor_data()
    print(factor_data)
    print("1.2 Load the calendar data")
    trade_days = get_calendar_data()
    print("1.3 Load the industry data")
    industry_code_df = get_stock_industry()
    print("1.4 Load the stock return")
    stock_return = get_stock_return()
    print("1.5 Merge the stock return and factor data")
    factor_neut = pd.merge(stock_return, factor_data, how="left", on=["code", "date"])

    drop_industry = ["801780", "801790"]
    index_code = "000905.WI"
    a = time.time()
    #################################################################################################### yb

    print("2. Running factor test")

    results = factor_pool_test_simple(
        factor_neut,
        trade_days,
        stock_return,
        industry_code_df,
        drop_industry,
        index_code,
        param,
    )
    # Here, results is a List[df] type returned after multiprocessing (mp) computation, where each df is a single-row DataFrame representing the backtest results of a factor's returns.
    result_df = pd.concat(results, axis=0)
    # result_df is a concatenated result.

    print("3. Generate and save the document")
    add_heading("Neutralized factor correlation", doc, level=1, seq=2)
    chosen_col = [col_ for col_ in factor_neut.columns if "neut" in col_]
    neut_corr = factor_neut[chosen_col].corr()
    add_heading_nonum(
        "Heatmap of factor correlations after neutralization", doc, level=2, seq=0
    )
    draw_thermodynamic_diagram(neut_corr, doc, "相关性热力图", inch=10.5)

    style = "Light List Accent 2"
    add_heading("Factor performance", doc, level=1, seq=2)
    barra_df = result_df[result_df["因子"]["因子名"].isin(param.barra_factor)]
    tech_df = result_df[result_df["因子"]["因子名"].isin(param.tech_facotr)]
    fin_df = result_df[result_df["因子"]["因子名"].isin(param.fin_factor)]
    add_heading_nonum("Barra factor performance", doc, level=2, seq=1)
    add_table(barra_df.T.reset_index().T, doc, size=7, style=style, width=8)
    add_heading_nonum("Technical factor performance", doc, level=2, seq=2)
    add_table(tech_df.T.reset_index().T, doc, size=7, style=style, width=8)
    add_heading_nonum("Fundamental factor performance", doc, level=2, seq=3)
    add_table(fin_df.T.reset_index().T, doc, size=7, style=style, width=8)
    # result_df.to_excel(os.path.join(param.pool_test, '因子监控.xlsx'))
    # result_df = pd.read_excel(os.path.join(param.pool_test, '因子监控.xlsx'))

    # 3、Factor screening results
    add_heading("Factor screening results", doc, level=1, seq=3)
    factor_success = pd.DataFrame(columns=result_df.columns)
    factor_fail = pd.DataFrame(columns=["因子", "分类", "未通过原因"])

    # 进行因子表现剔除
    for factor in param.factor_list:
        factor_df = result_df[result_df["因子"]["因子名"] == factor]
        factor_cate = factor_df["因子"]["分类"].iloc[0]
        factor_cate = "技术面"  ############################## 删
        IC_thresh = param.pool_thresh[factor_cate]["IC"]
        t_value_thresh = param.pool_thresh[factor_cate]["t_value"]
        direction_rate_thresh = param.pool_thresh[factor_cate]["direction_rate"]
        ARR_thresh = param.pool_thresh[factor_cate]["ARR"]
        win_rate_thresh = param.pool_thresh[factor_cate]["win_rate"]
        sharp_thresh = param.pool_thresh[factor_cate]["sharp"]

        factor_word = ""
        if float(factor_df["RankIC"]["均值"].iloc[0][:-1]) < IC_thresh * 100:
            factor_word += "因子RankIC的均值过低；"
        if float(factor_df["RankIC"]["t值"].iloc[0]) < t_value_thresh:
            factor_word += "因子RankIC的t值过低；"
        if (
            float(factor_df["RankIC"]["方向延续概率"].iloc[0][:-1])
            < direction_rate_thresh * 100
        ):
            factor_word += "因子RankIC的方向延续概率过低；"
        if float(factor_df["多空组合"]["年化收益"].iloc[0][:-1]) < ARR_thresh * 100:
            factor_word += "因子多空组合收益率过低；"
        if float(factor_df["多空组合"]["收益率t值"].iloc[0]) < t_value_thresh:
            factor_word += "因子多空收益率的t值过低；"
        if float(factor_df["多空组合"]["夏普比"].iloc[0]) < sharp_thresh:
            factor_word += "因子多空组合夏普比过低；"
        if float(factor_df["多空组合"]["胜率"].iloc[0][:-1]) < win_rate_thresh * 100:
            factor_word += "因子多空组合胜率过低；"

        if factor_word == "":  # 因子通过筛选
            factor_success = pd.concat([factor_success, factor_df], axis=0)
        else:
            factor_fail.loc[len(factor_fail)] = [factor, factor_cate, factor_word]

    print("4. Factor correlation elimination")
    success_list = list(factor_success.index)
    neut_corr_success = neut_corr.copy()
    col_name = [col_[:-5] for col_ in neut_corr_success.columns]
    neut_corr_success.columns = col_name
    neut_corr_success.index = col_name

    neut_corr_success = neut_corr_success[success_list].T[success_list]
    delete_factor = []
    for factor_ in success_list:
        factor_corr = neut_corr_success[factor_]
        try:
            factor_del_list = list(factor_corr[abs(factor_corr) > 0.9].index)
            factor_del_list.append(factor_)
            del_df = result_df[result_df.index.isin(factor_del_list)]
            del_df[("多空组合", "夏普比")] = del_df.apply(
                lambda x: float(x[("多空组合", "夏普比")]), axis=1
            )
            del_df.sort_values(
                by=[("多空组合", "夏普比")], ascending=False, inplace=True
            )  # sort by sharp ratio
            delete_factor += list(del_df.index)[1:]  # 只保存夏普最高的因子
        except:
            pass
    if len(delete_factor) != 0:  # Existence of highly correlated factors
        delete_factor = set(delete_factor)  # Deduplication
        for factor_ in delete_factor:
            factor_success = factor_success[factor_success.index != factor_]
            factor_cate = result_df[result_df["因子"]["因子名"] == factor_]["因子"][
                "分类"
            ].iloc[0]
            factor_word = "与已入选因子相关性过高"
            factor_fail.loc[len(factor_fail)] = [factor_, factor_cate, factor_word]

    add_heading("Selected factor performance", doc, level=3, seq=1)
    add_table(factor_success.T.reset_index().T, doc, size=7, style=style, width=8)
    add_heading("Reasons for the exclusion of unselected factors", doc, level=3, seq=2)
    add_table(factor_fail.T.reset_index().T, doc, size=7, style=style, width=8)

    print("5. Individual factor performance output")
    add_heading("Individual factor performance", doc, level=1, seq=4)
    num = 0
    for factor in param.factor_list:
        num += 1
        add_heading_nonum("%s factor performance" % factor, doc, level=2, seq=num)
        # 因子收益回测
        factor_result = result_df[result_df["因子"]["因子名"] == factor]
        add_heading("Return backtest", doc, level=3, seq=1)
        add_table(factor_result.T.reset_index().T, doc, size=7, style=style, width=8)
        # 因子相关性
        add_heading("Top 10 factor correlations", doc, level=3, seq=2)
        factor_corr = neut_corr[[factor + "_neut"]]
        factor_corr["abs"] = abs(factor_corr[factor + "_neut"])
        factor_corr.index = [factor_[:-5] for factor_ in factor_corr.index]
        factor_corr = factor_corr[factor_corr["abs"] != 1]  # 删除因子本身
        factor_corr.sort_values(by="abs", ascending=False, inplace=True)
        factor_corr = factor_corr.iloc[:10, :]  # 取前十相关因子
        factor_corr = round(factor_corr, 2)  # 保留两位小数
        del factor_corr["abs"]
        add_table(
            factor_corr.reset_index().T.reset_index(), doc, size=7, style=style, width=8
        )

        add_heading("Factor backtest curve", doc, level=3, seq=3)
        images = os.path.join(
            param.pool_test + "%s_净值曲线及回撤（多空收益）.png" % factor
        )
        doc_add_images(doc, images, inch=6)
        images = os.path.join(
            param.pool_test + "%s_净值曲线及回撤（相对收益）.png" % factor
        )
        doc_add_images(doc, images, inch=6)

    doc.save("FactorPoolScreen-v2.docx")  # 保存路径
    # doc.save(os.path.join(param.pool_test, '监控.docx'))  # 保存路径
