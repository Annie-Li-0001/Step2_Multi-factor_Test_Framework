from docx.shared import Pt
from docx.shared import Inches
from PIL import Image
import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 设置对象居中、对齐等。
from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
)  # 设置表格居中、表格行高


def doc_add_dataframe(doc, dataframe, fontsize):
    """

    :param doc: 文档对象
    :param dataframe:
    :param fontsize: 字体大小
    :return:
    """
    # 添加一个表格，style引入样式
    data = dataframe.copy()
    for i in data.columns:
        try:
            data[i] = data[i].round(decimals=4)
            data[i] = data[i].astype(str)
        except:
            data[i] = data[i].astype(str)
    row_num = data.shape[0]
    col_num = data.shape[1]
    doc_table = doc.add_table(rows=1, cols=col_num, style="Medium List 1 Accent 1")

    # 设置表头
    for i in range(col_num):
        doc_table.rows[0].cells[i].text = data.columns[i]

    # 添加数据，add_row()新增一行
    for i in range(row_num):
        new_row = doc_table.add_row()
        for j in range(col_num):
            new_row.cells[j].text = data.iloc[i, j]

    # 修改表格字体大小
    for row in doc_table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(fontsize)


def doc_add_images(doc, images, inch=3):
    try:
        doc.add_picture(images, width=Inches(inch))  # 添加图, 设置宽度
    except Exception:
        jpg_ima = Image.open(images)  # 打开图片
        jpg_ima.save("0.jpg")  # 保存新的图片
        doc.add_picture(images, width=Inches(inch))  # 添加图, 设置宽度


# 添加标题
def add_heading(title, doc, level, seq):
    title_num = {
        1: [
            "一、",
            "二、",
            "三、",
            "四、",
            "五、",
            "六、",
            "七、",
            "八、",
            "九、",
            "十、",
            "十一、",
            "十二、",
            "十三、",
            "十四、",
            "十五、",
        ],
        2: [
            "1. ",
            "2. ",
            "3. ",
            "4. ",
            "5. ",
            "6. ",
            "7. ",
            "8. ",
            "9. ",
            "10. ",
            "11. ",
            "12. ",
            "13. ",
            "14. ",
            "15. ",
            "16. ",
            "17. ",
            "18. ",
            "19. ",
            "20. ",
            "21. ",
            "22. ",
            "23. ",
            "24. ",
            "25. ",
            "26. ",
            "27. ",
            "28. ",
            "29. ",
            "30. ",
            "31. ",
            "32. ",
            "33. ",
            "34. ",
            "35. ",
            "36. ",
            "37. ",
            "38. ",
            "39. ",
            "40. ",
        ],
        3: [
            "(1) ",
            "(2) ",
            "(3) ",
            "(4) ",
            "(5) ",
            "(6) ",
            "(7) ",
            "(8) ",
            "(9) ",
            "(10) ",
        ],
    }
    if level == 0:
        doc.add_heading(title, level=level)
    else:
        doc.add_heading(title_num[level][seq - 1] + title, level=level)


# 添加表格
def add_table(table_df, doc, size=10, style="Table Grid", width=8):
    table = doc.add_table(
        rows=table_df.shape[0], cols=table_df.shape[1], style=style
    )  # Light Grid Accent 1
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
    for i in range(len(table_df)):
        for j in range(table_df.shape[1]):
            cell = table.cell(i, j)
            cell.width = Inches(width / table_df.shape[1])
            if i == 0 and j != 0 and table_df.iloc[i, j - 1] == table_df.iloc[i, j]:
                cell.merge(table.cell(i, j - 1))
            if pd.isna(table_df.iloc[i, j]):
                cell.text = ""
            else:
                cell.text = str(table_df.iloc[i, j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直居中
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中
            table.rows[i].height = Inches(0.25)
    table.style.font.size = Pt(size)
    table.autofit = False


def update_monitor_doc(IC_trends, return_trends, param, stock_pool):
    from docx import Document

    doc = Document()  # doc对象
    doc.add_heading(param.end_date + param.name + "因子监控", level=0)

    doc.add_heading("1 IC统计", level=1)  # 添加标题

    doc.add_heading("1.1 各市场近期IC表现", level=2)  # 添加标题
    doc_add_dataframe(doc, IC_trends)

    doc.add_heading("1.2 周度累计IC值", level=2)
    for i in stock_pool:
        images = param.path_monitor_results + i + "_周度累计IC.png"
        doc_add_images(doc, images)

    doc.add_heading("1.3 近半年月度IC值", level=2)
    for i in stock_pool:
        images = (
            param.path_monitor_results
            + param.end_date
            + "_近半年月度IC值("
            + i
            + ").png"
        )
        doc_add_images(doc, images, inch=5)

    doc.add_heading("2 收益统计", level=1)  # 添加文字

    doc.add_heading("2.1 各市场近期收益率表现", level=2)  # 添加文字
    doc_add_dataframe(doc, return_trends)

    doc.add_heading("2.2 多头累计收益", level=2)  # 添加文字
    for i in stock_pool:
        images = param.path_monitor_results + i + "_多头累计收益.png"
        doc_add_images(doc, images)

    doc.add_heading("2.3 近半年月度收益率", level=2)  # 添加文字
    for i in stock_pool:
        images = (
            param.path_monitor_results
            + param.end_date
            + "_近半年月度收益("
            + i
            + ").png"
        )
        doc_add_images(doc, images, inch=5)

    # doc.add_heading('3 持仓变动情况', level=1)
    # for factor in param.factors_name:
    #     for pool in stock_pool:
    #         w_hold_change = pd.read_csv(param.path_monitor_results+'持仓变动监控/'+factor+'/'+param.end_date+'_'+'w'+'_'+pool+'_'+factor+'_'+str(param.group_num)+'持仓变动情况.csv',index_col=0)
    #         m_hold_change = pd.read_csv(param.path_monitor_results+'持仓变动监控/'+factor+'/'+param.end_date+'_'+'m'+'_'+pool+'_'+factor+'_'+str(param.group_num)+'持仓变动情况.csv',index_col=0)
    #         doc.add_heading('3.1 周频：'+pool+' '+factor+'因子持仓变动情况', level=2)
    #         doc_add_dataframe(doc, w_hold_change)
    #         doc.add_heading('3.2 月频：'+pool+' '+factor+'因子持仓变动情况', level=2)
    #         doc_add_dataframe(doc, m_hold_change)
    doc.save(
        param.path_doc_results + param.end_date + param.name + "监控.docx"
    )  # 保存路径
    DESTINATIONS = ["sdlwqk@126.com"]
    title = param.name + "因子监控"
    content = "附件为" + param.name + "因子监控文档"
    # SendMail(DESTINATIONS, title, content).send([param.path_doc_results+param.end_date+param.name+'监控.docx'])
